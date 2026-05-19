# -*- coding: utf-8 -*-
"""
每周考勤表生成脚本
完全参照 qiandao.docx 格式
"""

import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime, timedelta
import os


def get_week_dates(year, week_num):
    """获取指定年份和周数的周一到周日日期"""
    jan1 = datetime(year, 1, 1)
    days_to_monday = (7 - jan1.weekday()) % 7
    if days_to_monday == 0 and jan1.weekday() != 0:
        days_to_monday = 7
    first_monday = jan1 + timedelta(days=days_to_monday)
    target_monday = first_monday + timedelta(weeks=week_num - 1)
    week_dates = []
    for i in range(7):
        week_dates.append(target_monday + timedelta(days=i))
    return week_dates


def get_week_range_from_dates(start_date, end_date):
    """从日期范围获取周数范围"""
    start_monday = start_date - timedelta(days=start_date.weekday())
    year = start_date.year
    weeks = []
    current_monday = start_monday
    while current_monday <= end_date:
        week_num = current_monday.isocalendar()[1]
        if week_num not in weeks:
            weeks.append(week_num)
        current_monday += timedelta(days=7)
    return year, min(weeks), max(weeks)


def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_diagonal_border(cell):
    """设置单元格斜线（从左上到右下）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:tl2br w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_run_font(run, name='宋体', size=Pt(12), bold=None, color=None):
    """统一设置字体"""
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def set_cell_vertical_center(cell):
    """设置单元格垂直居中"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def create_attendance_table(doc, year, week_num, employees, title_text="每周考勤表"):
    """创建单周考勤表 - 完全参照原文件格式"""
    week_dates = get_week_dates(year, week_num)

    # 每页最多13行数据行（考虑标题、日期范围、表头、行间距等因素）
    max_data_rows_per_page = 13
    num_cols = 10

    # 将员工分页处理
    pages = []
    for i in range(0, len(employees), max_data_rows_per_page):
        pages.append(employees[i:i + max_data_rows_per_page])

    # 如果没有员工，至少生成一页空白
    if not pages:
        pages.append([])

    for page_idx, page_employees in enumerate(pages):
        # 第二页开始添加分页符
        if page_idx > 0:
            doc.add_page_break()

        # 创建标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        set_run_font(run, size=Pt(18), bold=True)

        # 创建日期范围
        date_range = doc.add_paragraph()
        date_range.alignment = WD_ALIGN_PARAGRAPH.CENTER
        start_date = week_dates[0].strftime("%Y年%m月%d日")
        end_date = week_dates[6].strftime("%Y年%m月%d日")
        run = date_range.add_run(f"日期范围：   {start_date}                            至     {end_date}")
        set_run_font(run, size=Pt(10))

        # 计算本页数据行数：员工数或最大行数，取较大者
        num_data_rows = max(len(page_employees), max_data_rows_per_page)
        num_rows = num_data_rows + 2  # +2行表头

        # 创建表格
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 设置单元格宽度
        for row in table.rows:
            for i in range(num_cols):
                row.cells[i].width = Cm(2.45)

        # 设置行高：表头每行0.84cm，数据行0.75cm
        for i, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            height_val = "477" if i < 2 else "425"
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_val}" w:hRule="atLeast"/>')
            trPr.append(trHeight)

        # === 表头第一行 ===
        # 职位（合并两行，不加粗，垂直居中）
        cell = table.cell(0, 0)
        cell.text = "职位"
        set_cell_vertical_center(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=Pt(12))

        # 姓名单元格（合并两行，带斜线）
        cell = table.cell(0, 1)
        for p in cell.paragraphs:
            p.clear()
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run("日期")
        set_run_font(run, size=Pt(12))
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        run = p3.add_run("姓名")
        set_run_font(run, size=Pt(12))
        set_cell_diagonal_border(cell)

        # 日期列 - 蓝色字体，不加粗，垂直居中
        for i, date in enumerate(week_dates):
            cell = table.cell(0, i + 2)
            cell.text = date.strftime("%m月%d日")
            set_cell_vertical_center(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=Pt(12), color=RGBColor(0, 0, 255))

        # 备注（合并两行，不加粗，垂直居中）
        cell = table.cell(0, 9)
        cell.text = "备注"
        set_cell_vertical_center(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=Pt(12))

        # === 表头第二行 ===
        cell = table.cell(1, 0)
        cell.text = ""

        cell = table.cell(1, 1)
        cell.text = ""

        # 星期列 - 蓝色字体，不加粗，垂直居中，周末绿色背景
        week_days = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
        for i, day in enumerate(week_days):
            cell = table.cell(1, i + 2)
            cell.text = day
            set_cell_vertical_center(cell)
            if i in [5, 6]:
                set_cell_shading(cell, "03D96A")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=Pt(12), color=RGBColor(0, 0, 255))

        cell = table.cell(1, 9)
        cell.text = ""

        # 合并单元格：职位、姓名、备注
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(1, 1))
        table.cell(0, 9).merge(table.cell(1, 9))

        # === 数据行 ===
        for row_idx in range(num_data_rows):
            data_row = row_idx + 2

            if row_idx < len(page_employees):
                position, name = page_employees[row_idx]
            else:
                position, name = "", ""

            # 职位（垂直居中）
            cell = table.cell(data_row, 0)
            cell.text = position
            set_cell_vertical_center(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=Pt(12))

            # 姓名（垂直居中）
            cell = table.cell(data_row, 1)
            cell.text = name
            set_cell_vertical_center(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=Pt(12))

            # 周一到周日留空（垂直居中）
            for col in range(2, 9):
                cell = table.cell(data_row, col)
                cell.text = ""
                set_cell_vertical_center(cell)

            # 备注留空（垂直居中）
            cell = table.cell(data_row, 9)
            cell.text = ""
            set_cell_vertical_center(cell)

    return doc


def generate_attendance_sheets(year, start_week, end_week, employees, title_text="每周考勤表"):
    """生成连续的每周考勤表"""
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(3.17)
    section.bottom_margin = Cm(3.17)

    for week_num in range(start_week, end_week + 1):
        create_attendance_table(doc, year, week_num, employees, title_text)
        if week_num < end_week:
            doc.add_page_break()

    return doc


class AttendanceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("每周考勤表生成器")
        self.root.geometry("600x600")
        self.root.resizable(False, False)
        self.employees = []
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # 标题设置
        title_frame = ttk.LabelFrame(main_frame, text="标题设置", padding="10")
        title_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))

        ttk.Label(title_frame, text="标题：").grid(row=0, column=0, padx=(0, 5))
        self.title_var = tk.StringVar(value="每周考勤表")
        ttk.Entry(title_frame, textvariable=self.title_var, width=40).grid(row=0, column=1)

        # 日期设置
        date_frame = ttk.LabelFrame(main_frame, text="日期范围", padding="10")
        date_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))

        ttk.Label(date_frame, text="从：").grid(row=0, column=0, padx=(0, 5))
        self.start_date_var = tk.StringVar(value="2026-05-18")
        ttk.Entry(date_frame, textvariable=self.start_date_var, width=12).grid(row=0, column=1, padx=(0, 20))

        ttk.Label(date_frame, text="到：").grid(row=0, column=2, padx=(0, 5))
        self.end_date_var = tk.StringVar(value="2026-06-15")
        ttk.Entry(date_frame, textvariable=self.end_date_var, width=12).grid(row=0, column=3)

        ttk.Label(date_frame, text="格式：YYYY-MM-DD").grid(row=1, column=0, columnspan=4, pady=(5, 0))

        # 员工设置
        emp_frame = ttk.LabelFrame(main_frame, text="员工信息", padding="10")
        emp_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))

        add_frame = ttk.Frame(emp_frame)
        add_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))

        ttk.Label(add_frame, text="职位：").grid(row=0, column=0, padx=(0, 5))
        self.position_var = tk.StringVar()
        ttk.Entry(add_frame, textvariable=self.position_var, width=15).grid(row=0, column=1, padx=(0, 10))

        ttk.Label(add_frame, text="姓名：").grid(row=0, column=2, padx=(0, 5))
        self.name_var = tk.StringVar()
        ttk.Entry(add_frame, textvariable=self.name_var, width=15).grid(row=0, column=3, padx=(0, 10))

        ttk.Button(add_frame, text="添加", command=self.add_employee).grid(row=0, column=4)

        list_frame = ttk.Frame(emp_frame)
        list_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E))

        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
        self.emp_listbox = tk.Listbox(list_frame, height=10, yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.emp_listbox.yview)
        self.emp_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        btn_frame = ttk.Frame(emp_frame)
        btn_frame.grid(row=2, column=0, columnspan=3, pady=(10, 0))
        ttk.Button(btn_frame, text="删除选中", command=self.delete_employee).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="清空", command=self.clear_employees).pack(side=tk.LEFT)

        ttk.Button(main_frame, text="生成考勤表", command=self.generate).grid(row=3, column=0, columnspan=2, pady=20)

    def add_employee(self):
        position = self.position_var.get().strip()
        name = self.name_var.get().strip()
        if not position or not name:
            messagebox.showwarning("提示", "请输入职位和姓名")
            return
        self.employees.append((position, name))
        self.emp_listbox.insert(tk.END, f"{position} - {name}")
        self.position_var.set("")
        self.name_var.set("")

    def delete_employee(self):
        selection = self.emp_listbox.curselection()
        if selection:
            index = selection[0]
            self.emp_listbox.delete(index)
            self.employees.pop(index)

    def clear_employees(self):
        self.emp_listbox.delete(0, tk.END)
        self.employees.clear()

    def generate(self):
        if not self.employees:
            messagebox.showwarning("提示", "请添加至少一名员工")
            return

        try:
            start_str = self.start_date_var.get().strip()
            end_str = self.end_date_var.get().strip()

            start = datetime.strptime(start_str, "%Y-%m-%d")
            end = datetime.strptime(end_str, "%Y-%m-%d")

            if start > end:
                messagebox.showerror("错误", "开始日期不能晚于结束日期")
                return

            year, start_week, end_week = get_week_range_from_dates(start, end)

            title_text = self.title_var.get().strip()
            if not title_text:
                title_text = "每周考勤表"

            doc = generate_attendance_sheets(year, start_week, end_week, self.employees, title_text)

            if start_week == end_week:
                filename = f"{title_text}_{year}年第{start_week}周.docx"
            else:
                filename = f"{title_text}_{year}年第{start_week}-{end_week}周.docx"

            doc.save(filename)
            messagebox.showinfo("成功", f"考勤表已生成：\n{filename}\n\n保存位置：{os.getcwd()}")
        except ValueError:
            messagebox.showerror("错误", "日期格式错误，请使用 YYYY-MM-DD 格式")
        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{str(e)}")


def main():
    root = tk.Tk()
    app = AttendanceApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
